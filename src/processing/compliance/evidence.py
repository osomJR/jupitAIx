from __future__ import annotations

"""
Evidence extraction utilities for compliance evaluation.

Design goals:
- deterministic, text-first evidence discovery
- preserve source_document_index always
- preserve page_number for PDFs when possible
- preserve section_label when inferable from nearby text
- return schema-aligned EvidenceReference objects
- avoid inventing evidence when the source text is weak or unavailable
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional, Sequence
import re

import docx  # type: ignore
import fitz  # PyMuPDF

try:
    from src.schema import DocumentInputFormat, DocumentPayload, DocumentSetPayload, EvidenceReference
except ImportError:  # pragma: no cover
    from schema import DocumentInputFormat, DocumentPayload, DocumentSetPayload, EvidenceReference

try:
    from src.extraction import extract_text_by_format
except ImportError:  # pragma: no cover
    from extraction import extract_text_by_format

DEFAULT_EXCERPT_WINDOW = 160
DEFAULT_MAX_MATCHES_PER_SIGNAL = 5


@dataclass(frozen=True)
class PageText:
    page_number: int
    text: str


@dataclass(frozen=True)
class EvidenceDocument:
    source_document_index: int
    input_format: DocumentInputFormat
    source_reference: Optional[str]
    text: str
    pages: tuple[PageText, ...]


@dataclass(frozen=True)
class SignalMatch:
    source_document_index: int
    signal: str
    start: int
    end: int
    page_number: Optional[int]
    section_label: Optional[str]
    locator_text: str
    excerpt: str

    def to_reference(self) -> EvidenceReference:
        return EvidenceReference(
            source_document_index=self.source_document_index,
            page_number=self.page_number,
            section_label=self.section_label,
            locator_text=self.locator_text,
            excerpt=self.excerpt,
        )


@dataclass(frozen=True)
class EvidenceSearchResult:
    signal: str
    matches: tuple[SignalMatch, ...]


class EvidenceExtractionError(RuntimeError):
    """Raised when evidence extraction cannot proceed safely."""


def build_evidence_documents(
    input_artifact: DocumentPayload | DocumentSetPayload,
) -> list[EvidenceDocument]:
    """
    Build a normalized evidence corpus from a single document or document set.

    Notes:
    - Uses payload text when already available.
    - Falls back to reopening the source file only when the supplied file reference
      resolves to a readable local path.
    - Preserves page text for PDFs when a readable PDF path is available.
    - Uses extraction.extract_text_by_format as a best-effort OCR fallback when
      local payload text is absent.
    """
    documents = [input_artifact] if isinstance(input_artifact, DocumentPayload) else list(input_artifact.documents)
    corpus: list[EvidenceDocument] = []

    for index, document in enumerate(documents):
        source_path = _resolve_source_path(document.filename)
        input_format = document.metadata.input_format

        pages: list[PageText] = []
        text = _normalize_text(document.text)

        if input_format == DocumentInputFormat.pdf and source_path is not None:
            pages = _extract_pdf_pages(source_path)
            if not text:
                text = "\n\n".join(page.text for page in pages if page.text.strip()).strip()
            if not text:
                text = _extract_text_with_fallback(source_path, input_format)
        elif not text and source_path is not None:
            text = _extract_text_with_fallback(source_path, input_format)

        corpus.append(
            EvidenceDocument(
                source_document_index=index,
                input_format=input_format,
                source_reference=str(source_path) if source_path is not None else document.filename,
                text=text,
                pages=tuple(page for page in pages if page.text.strip()),
            )
        )

    return corpus


def collect_evidence_references(
    documents: Sequence[EvidenceDocument],
    *,
    signals: Sequence[str],
    search_mode: str = "substring",
    case_sensitive: bool = False,
    excerpt_window: int = DEFAULT_EXCERPT_WINDOW,
    max_matches_per_signal: int = DEFAULT_MAX_MATCHES_PER_SIGNAL,
) -> list[EvidenceReference]:
    references: list[EvidenceReference] = []
    for signal in signals:
        result = search_signal(
            documents,
            signal=signal,
            search_mode=search_mode,
            case_sensitive=case_sensitive,
            excerpt_window=excerpt_window,
            max_matches=max_matches_per_signal,
        )
        references.extend(match.to_reference() for match in result.matches)
    return _dedupe_references(references)


def search_signal(
    documents: Sequence[EvidenceDocument],
    *,
    signal: str,
    search_mode: str = "substring",
    case_sensitive: bool = False,
    excerpt_window: int = DEFAULT_EXCERPT_WINDOW,
    max_matches: int = DEFAULT_MAX_MATCHES_PER_SIGNAL,
) -> EvidenceSearchResult:
    normalized_signal = signal.strip()
    if not normalized_signal:
        return EvidenceSearchResult(signal=signal, matches=())

    matches: list[SignalMatch] = []
    for document in documents:
        if not document.text:
            continue

        if document.pages:
            for page in document.pages:
                page_matches = _find_matches_in_text(
                    page.text,
                    normalized_signal,
                    search_mode=search_mode,
                    case_sensitive=case_sensitive,
                    excerpt_window=excerpt_window,
                    max_matches=max_matches,
                )
                for start, end, locator_text, excerpt in page_matches:
                    matches.append(
                        SignalMatch(
                            source_document_index=document.source_document_index,
                            signal=normalized_signal,
                            start=start,
                            end=end,
                            page_number=page.page_number,
                            section_label=_infer_section_label(page.text, start),
                            locator_text=locator_text,
                            excerpt=excerpt,
                        )
                    )
                    if len(matches) >= max_matches:
                        return EvidenceSearchResult(signal=normalized_signal, matches=tuple(matches))
        else:
            document_matches = _find_matches_in_text(
                document.text,
                normalized_signal,
                search_mode=search_mode,
                case_sensitive=case_sensitive,
                excerpt_window=excerpt_window,
                max_matches=max_matches,
            )
            for start, end, locator_text, excerpt in document_matches:
                matches.append(
                    SignalMatch(
                        source_document_index=document.source_document_index,
                        signal=normalized_signal,
                        start=start,
                        end=end,
                        page_number=None,
                        section_label=_infer_section_label(document.text, start),
                        locator_text=locator_text,
                        excerpt=excerpt,
                    )
                )
                if len(matches) >= max_matches:
                    return EvidenceSearchResult(signal=normalized_signal, matches=tuple(matches))

    return EvidenceSearchResult(signal=normalized_signal, matches=tuple(matches))


def get_source_reference(
    input_artifact: DocumentPayload | DocumentSetPayload,
    *,
    source_document_index: int,
) -> Optional[str]:
    documents = [input_artifact] if isinstance(input_artifact, DocumentPayload) else list(input_artifact.documents)
    if source_document_index < 0 or source_document_index >= len(documents):
        return None
    return documents[source_document_index].filename


def _resolve_source_path(source_reference: Optional[str]) -> Optional[Path]:
    if not source_reference:
        return None
    candidate = Path(source_reference)
    return candidate if candidate.exists() else None


def _extract_pdf_pages(source_path: Path) -> list[PageText]:
    pages: list[PageText] = []
    with fitz.open(source_path) as pdf:
        for index, page in enumerate(pdf, start=1):
            page_text = page.get_text().strip()
            pages.append(PageText(page_number=index, text=page_text))
    return pages


def _extract_text_with_fallback(source_path: Path, input_format: DocumentInputFormat) -> str:
    if input_format == DocumentInputFormat.txt:
        return source_path.read_text(encoding="utf-8").strip()
    if input_format == DocumentInputFormat.docx:
        document = docx.Document(source_path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
    extracted_text, _ = extract_text_by_format(source_path, input_format)
    return extracted_text.strip()


def _find_matches_in_text(
    text: str,
    signal: str,
    *,
    search_mode: str,
    case_sensitive: bool,
    excerpt_window: int,
    max_matches: int,
) -> list[tuple[int, int, str, str]]:
    if not text:
        return []

    mode = search_mode.lower().strip()
    flags = 0 if case_sensitive else re.IGNORECASE

    if mode == "regex":
        pattern = re.compile(signal, flags)
        found: list[tuple[int, int, str, str]] = []
        for match in pattern.finditer(text):
            start, end = match.span()
            found.append((start, end, text[start:end].strip(), _slice_excerpt(text, start, end, excerpt_window)))
            if len(found) >= max_matches:
                break
        return found

    haystack = text if case_sensitive else text.lower()
    needle = signal if case_sensitive else signal.lower()
    found = []
    start_index = 0
    while len(found) < max_matches:
        start = haystack.find(needle, start_index)
        if start < 0:
            break
        end = start + len(signal)
        found.append((start, end, text[start:end].strip(), _slice_excerpt(text, start, end, excerpt_window)))
        start_index = end
    return found


def _slice_excerpt(text: str, start: int, end: int, excerpt_window: int) -> str:
    left = max(0, start - excerpt_window)
    right = min(len(text), end + excerpt_window)
    excerpt = text[left:right].strip()
    excerpt = re.sub(r"\s+", " ", excerpt)
    if len(excerpt) > 400:
        excerpt = excerpt[:397] + "..."
    return excerpt


def _infer_section_label(text: str, position: int) -> Optional[str]:
    if not text:
        return None

    prefix = text[:position]
    lines = [line.strip() for line in prefix.splitlines() if line.strip()]
    for candidate in reversed(lines[-8:]):
        if _looks_like_heading(candidate):
            return candidate[:200]
    return None


def _looks_like_heading(value: str) -> bool:
    if not value:
        return False
    compact = re.sub(r"\s+", " ", value).strip()
    if len(compact) > 120:
        return False
    if compact.endswith(":"):
        return True
    if compact.isupper() and len(compact) <= 80:
        return True
    words = compact.split()
    titleish = sum(1 for word in words if word[:1].isupper())
    return len(words) <= 10 and titleish >= max(1, len(words) // 2)


def _normalize_text(value: Optional[str]) -> str:
    return value.strip() if isinstance(value, str) else ""


def _dedupe_references(references: Iterable[EvidenceReference]) -> list[EvidenceReference]:
    seen: set[tuple[int, Optional[int], Optional[str], Optional[str], Optional[str]]] = set()
    deduped: list[EvidenceReference] = []
    for reference in references:
        key = (
            reference.source_document_index,
            reference.page_number,
            reference.section_label,
            reference.locator_text,
            reference.excerpt,
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(reference)
    return deduped


__all__ = [
    "DEFAULT_EXCERPT_WINDOW",
    "DEFAULT_MAX_MATCHES_PER_SIGNAL",
    "EvidenceDocument",
    "EvidenceExtractionError",
    "EvidenceSearchResult",
    "PageText",
    "SignalMatch",
    "build_evidence_documents",
    "collect_evidence_references",
    "get_source_reference",
    "search_signal",
]
