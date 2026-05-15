from __future__ import annotations
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Mapping, Optional, Sequence
import re
import docx
import fitz  
import pytesseract
from PIL import Image, ImageFilter

from src.extraction import OCR_CONFIG, preprocess_for_ocr, resolve_ocr_lang
from src.processing.data_protection.client import (
    DEFAULT_DLP_LOCATION,
    DEFAULT_MIN_LIKELIHOOD,
    DetectionCandidate,
    GoogleSDPClient,
    TextFinding,
    build_google_sdp_client,
    inspect_sensitive_text,
    merge_overlapping_findings,
    preview_candidates_from_text,
)
from src.schema import (
    AnalyzerRequest,
    AnalyzerResponse,
    DataMaskingRequest,
    DocumentInputFormat,
    DocumentPayload,
    FeatureType,
    HumanReviewRequirement,
    SensitiveDataType,
)
from src.validation import build_document_file_result, validate_analyzer_request, validate_analyzer_response

DEFAULT_PDF_RENDER_SCALE = 2.0

_IMAGE_OUTPUT_MAP: dict[DocumentInputFormat, str] = {
    DocumentInputFormat.jpg: "JPEG",
    DocumentInputFormat.jpeg: "JPEG",
    DocumentInputFormat.png: "PNG",
}

_FILE_OUTPUT_MAP = {
    DocumentInputFormat.pdf: "pdf",
    DocumentInputFormat.docx: "docx",
    DocumentInputFormat.jpg: "jpg",
    DocumentInputFormat.jpeg: "jpeg",
    DocumentInputFormat.png: "png",
}

_SUPPORTED_INPUTS = {
    DocumentInputFormat.pdf,
    DocumentInputFormat.docx,
    DocumentInputFormat.jpg,
    DocumentInputFormat.jpeg,
    DocumentInputFormat.png,
}

_CUSTOM_MASK_LABEL = "custom_mask"


@dataclass(frozen=True)
class OCRWord:
    text: str
    bbox: tuple[int, int, int, int]
    start: int
    end: int


@dataclass(frozen=True)
class RunSpan:
    start: int
    end: int
    text: str
    rpr_xml: Any


@dataclass(frozen=True)
class PDFWord:
    start: int
    end: int
    text: str
    rect: fitz.Rect
    block_no: int
    line_no: int
    word_no: int


def _ocr_words_from_image(image: Image.Image, *, ocr_lang: Optional[str] = None) -> tuple[str, list[OCRWord]]:
    processed = preprocess_for_ocr(image.convert("RGB"))
    data = pytesseract.image_to_data(
        processed,
        lang=ocr_lang or resolve_ocr_lang(),
        config=OCR_CONFIG,
        output_type=pytesseract.Output.DICT,
    )

    words: list[OCRWord] = []
    text_parts: list[str] = []
    cursor = 0
    texts = data.get("text", [])
    lefts = data.get("left", [])
    tops = data.get("top", [])
    widths = data.get("width", [])
    heights = data.get("height", [])

    for idx, raw in enumerate(texts):
        token = (raw or "").strip()
        if not token:
            continue
        if text_parts:
            text_parts.append(" ")
            cursor += 1
        start = cursor
        text_parts.append(token)
        cursor += len(token)
        end = cursor
        bbox = (
            int(lefts[idx]),
            int(tops[idx]),
            int(lefts[idx]) + int(widths[idx]),
            int(tops[idx]) + int(heights[idx]),
        )
        words.append(OCRWord(text=token, bbox=bbox, start=start, end=end))
    return "".join(text_parts), words


def _normalize_text_for_compare(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip()).casefold()


def _clean_custom_mask_items(values: Optional[Sequence[str]] = None) -> list[str]:
    cleaned: list[str] = []
    seen: set[str] = set()

    for value in values or ():
        text = str(value).strip()
        if not text:
            continue

        key = _normalize_text_for_compare(text)
        if key in seen:
            continue

        cleaned.append(text)
        seen.add(key)

    return cleaned


def _custom_mask_findings(text: str, values: Optional[Sequence[str]] = None) -> list[TextFinding]:
    findings: list[TextFinding] = []

    for value in _clean_custom_mask_items(values):
        pattern = re.compile(re.escape(value), flags=re.IGNORECASE)
        for match in pattern.finditer(text):
            quote = text[match.start():match.end()]
            if not quote.strip():
                continue
            findings.append(
                TextFinding(
                    start=match.start(),
                    end=match.end(),
                    quote=quote,
                    label=_CUSTOM_MASK_LABEL,
                    source="custom_mask",
                )
            )

    return findings


def _mask_findings_for_text(
    *,
    sdp: GoogleSDPClient,
    text: str,
    payload: DataMaskingRequest,
    custom_redactions: Optional[Sequence[str]] = None,
) -> list[TextFinding]:
    sensitive_findings = inspect_sensitive_text(
        sdp=sdp,
        text=text,
        targets=payload.target_data,
        review_exclusions=payload.review_exclusions,
    )
    return merge_overlapping_findings(
        [*sensitive_findings, *_custom_mask_findings(text, custom_redactions)],
        original_text=text,
    )


def _boxes_for_text_spans(findings: Sequence[TextFinding], words: Sequence[OCRWord]) -> list[tuple[int, int, int, int]]:
    boxes: list[tuple[int, int, int, int]] = []
    for finding in findings:
        matched = [w for w in words if not (w.end <= finding.start or w.start >= finding.end)]
        if not matched:
            needle = _normalize_text_for_compare(finding.quote)
            for i in range(len(words)):
                buf = []
                for j in range(i, min(len(words), i + 12)):
                    buf.append(words[j])
                    candidate = _normalize_text_for_compare(" ".join(item.text for item in buf))
                    if candidate == needle:
                        matched = buf
                        break
                if matched:
                    break
        if not matched:
            continue
        x0 = min(w.bbox[0] for w in matched)
        y0 = min(w.bbox[1] for w in matched)
        x1 = max(w.bbox[2] for w in matched)
        y1 = max(w.bbox[3] for w in matched)
        boxes.append((x0, y0, x1, y1))

    unique: list[tuple[int, int, int, int]] = []
    seen: set[tuple[int, int, int, int]] = set()
    for box in boxes:
        if box not in seen:
            unique.append(box)
            seen.add(box)
    return unique


def _iter_table_paragraphs(table: Any):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_document_paragraphs(document: docx.Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _output_path(source_path: Path, output_dir: Path, suffix: str = "_masked") -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{source_path.stem}{suffix}{source_path.suffix.lower()}"


def _ensure_mask_request(request: AnalyzerRequest) -> tuple[DocumentPayload, DataMaskingRequest]:
    if request.action != FeatureType.data_mask:
        raise ValueError("data_mask.py only handles action='data_mask'.")
    if not isinstance(request.input, DocumentPayload):
        raise ValueError("data_mask requires DocumentPayload input.")
    if not isinstance(request.payload, DataMaskingRequest):
        raise ValueError("data_mask requires DataMaskingRequest payload.")
    if request.input.metadata.input_format not in _SUPPORTED_INPUTS:
        raise ValueError("data_mask only supports pdf, docx, jpg, jpeg, png.")
    if request.policy.structure_preservation is not True:
        raise ValueError("structure_preservation must be True for data_mask.")
    return request.input, request.payload


def _build_masking_response(*, request: AnalyzerRequest, output_path: Path) -> AnalyzerResponse:
    input_payload, _ = _ensure_mask_request(request)
    result = build_document_file_result(
        filename=output_path.name,
        output_format=_FILE_OUTPUT_MAP[input_payload.metadata.input_format],
        file_size_mb=round(output_path.stat().st_size / (1024 * 1024), 4),
        algorithm_version="google-sdp-data-mask-v1",
    )
    response = AnalyzerResponse(
        action=FeatureType.data_mask,
        input_format=input_payload.metadata.input_format,
        policy=request.policy,
        system_language=request.system_language,
        result=result,
        human_review=HumanReviewRequirement(),
    )
    return validate_analyzer_response(response, request=request)


def _mask_digits(value: str, *, keep_last: int = 4, mask_char: str = "X") -> str:
    digits = [i for i, char in enumerate(value) if char.isdigit()]
    if not digits:
        return value
    to_keep = set(digits[-keep_last:]) if keep_last > 0 else set()
    chars = list(value)
    for idx in digits:
        if idx not in to_keep:
            chars[idx] = mask_char
    return "".join(chars)


def _mask_alpha(value: str, *, keep_first: int = 1, mask_char: str = "X") -> str:
    chars = list(value)
    alpha_indices = [i for i, char in enumerate(chars) if char.isalpha()]
    to_keep = set(alpha_indices[:keep_first]) if keep_first > 0 else set()
    for idx in alpha_indices:
        if idx not in to_keep:
            chars[idx] = mask_char
    return "".join(chars)


def _mask_name(value: str) -> str:
    parts = value.split()
    masked = []
    for part in parts:
        masked.append(part[0] + ("X" * (len(part) - 1)) if len(part) > 1 else "X")
    return " ".join(masked) if masked else "X"


def _mask_email(value: str) -> str:
    if "@" not in value:
        return _mask_alpha(value, keep_first=1)
    local, domain = value.split("@", 1)
    if "." in domain:
        head, _, tld = domain.rpartition(".")
        masked_domain = (head[:1] + "X" * max(1, len(head) - 1)) if head else "X"
        return f"{local[:1]}{'X' * max(1, len(local) - 1)}@{masked_domain}.{tld}"
    return f"{local[:1]}{'X' * max(1, len(local) - 1)}@{'X' * max(1, len(domain))}"


def mask_value_by_target(label: str, quote: str) -> str:
    label = (label or "").lower()
    if label == _CUSTOM_MASK_LABEL:
        return "".join(char if char.isspace() else "X" for char in quote)
    if label in {SensitiveDataType.name.value, "person_name"}:
        return _mask_name(quote)
    if label in {SensitiveDataType.email_address.value, "email_address"}:
        return _mask_email(quote)
    if label in {SensitiveDataType.phone_number.value, "phone_number"}:
        return _mask_digits(quote, keep_last=2)
    if label in {SensitiveDataType.card_number.value, "credit_card_number", SensitiveDataType.account_number.value}:
        return _mask_digits(quote, keep_last=4)
    if label in {
        SensitiveDataType.national_id.value,
        SensitiveDataType.tax_id.value,
        SensitiveDataType.passport_number.value,
        "passport",
    }:
        return _mask_digits(_mask_alpha(quote, keep_first=1), keep_last=2)
    if label in {SensitiveDataType.contact_address.value, "street_address"}:
        stripped = quote.strip()
        return stripped[:6] + ("X" * max(0, len(stripped) - 6)) if len(stripped) > 6 else "X" * len(stripped)
    if label == SensitiveDataType.date_of_birth.value:
        chars = list(quote)
        digits = [i for i, char in enumerate(chars) if char.isdigit()]
        keep = set(digits[-2:])
        for idx in digits:
            if idx not in keep:
                chars[idx] = "X"
        return "".join(chars)
    if label == SensitiveDataType.age.value:
        return _mask_digits(quote, keep_last=0)
    if label == SensitiveDataType.signature.value:
        return "[SIGNATURE_MASKED]"
    return _mask_alpha(_mask_digits(quote, keep_last=2), keep_first=1)


def _same_length_mask_value(label: str, quote: str) -> str:
    masked = mask_value_by_target(label, quote)
    if len(masked) == len(quote):
        return masked
    if len(masked) < len(quote):
        return masked + ("X" * (len(quote) - len(masked)))
    return masked[: len(quote)]


def mask_text_with_findings(text: str, findings: Sequence[TextFinding]) -> str:
    result = text
    for finding in sorted(findings, key=lambda item: item.start, reverse=True):
        replacement = _same_length_mask_value(finding.label, finding.quote)
        result = result[:finding.start] + replacement + result[finding.end:]
    return result


def pixelate_boxes(
    image: Image.Image,
    boxes: Sequence[tuple[int, int, int, int]],
    *,
    blur_radius: int = 3,
) -> Image.Image:
    out = image.convert("RGB").copy()
    for x0, y0, x1, y1 in boxes:
        region = out.crop((x0, y0, x1, y1))
        width = max(1, x1 - x0)
        height = max(1, y1 - y0)
        reduced = region.resize((max(1, width // 10), max(1, height // 10)))
        expanded = reduced.resize((width, height))
        expanded = expanded.filter(ImageFilter.GaussianBlur(radius=blur_radius))
        out.paste(expanded, (x0, y0, x1, y1))
    return out


def _resolve_sdp(
    *,
    sdp: GoogleSDPClient | None = None,
    project_id: str | None = None,
    location: str = DEFAULT_DLP_LOCATION,
    min_likelihood: str = DEFAULT_MIN_LIKELIHOOD,
    client: Any | None = None,
) -> GoogleSDPClient:
    if sdp is not None:
        return sdp
    if not project_id:
        raise ValueError("Provide either sdp or project_id.")
    return build_google_sdp_client(
        project_id=project_id,
        location=location,
        min_likelihood=min_likelihood,
        client=client,
    )


def preview_data_mask_candidates(
    request: AnalyzerRequest | Mapping[str, Any],
    *,
    sdp: GoogleSDPClient | None = None,
    project_id: str | None = None,
    location: str = DEFAULT_DLP_LOCATION,
    min_likelihood: str = DEFAULT_MIN_LIKELIHOOD,
    client: Any | None = None,
    custom_redactions: Optional[Sequence[str]] = None,
) -> list[DetectionCandidate]:
    req = validate_analyzer_request(request)
    input_payload, payload = _ensure_mask_request(req)
    if not input_payload.text:
        raise ValueError(
            "preview_data_mask_candidates requires extracted document text. "
            "Build the request with extraction.py before previewing."
        )
    resolved = _resolve_sdp(
        sdp=sdp,
        project_id=project_id,
        location=location,
        min_likelihood=min_likelihood,
        client=client,
    )
    return preview_candidates_from_text(
        sdp=resolved,
        text=input_payload.text,
        targets=payload.target_data,
        review_exclusions=payload.review_exclusions,
        min_likelihood=min_likelihood,
    )


def _paragraph_run_spans(paragraph: Any) -> tuple[str, list[RunSpan]]:
    spans: list[RunSpan] = []
    parts: list[str] = []
    cursor = 0

    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        start = cursor
        end = start + len(text)
        parts.append(text)
        spans.append(
            RunSpan(
                start=start,
                end=end,
                text=text,
                rpr_xml=deepcopy(run._r.rPr) if run._r.rPr is not None else None,
            )
        )
        cursor = end

    return "".join(parts), spans


def _copy_rpr(dst_run: Any, rpr_xml: Any) -> None:
    if dst_run._r.rPr is not None:
        dst_run._r.remove(dst_run._r.rPr)
    if rpr_xml is not None:
        dst_run._r.insert(0, deepcopy(rpr_xml))


def _rewrite_paragraph_from_fragments(paragraph: Any, fragments: list[tuple[Any, str]]) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    for rpr_xml, text in fragments:
        if not text:
            continue
        new_run = paragraph.add_run()
        _copy_rpr(new_run, rpr_xml)
        new_run.text = text


def _run_span_for_offset(spans: Sequence[RunSpan], offset: int) -> RunSpan:
    for span in spans:
        if span.start <= offset < span.end:
            return span
    return spans[-1]


def mask_paragraph_runs(paragraph: Any, findings: Sequence[TextFinding]) -> None:
    full_text, spans = _paragraph_run_spans(paragraph)
    if not full_text or not spans or not findings:
        return

    merged = merge_overlapping_findings(findings, original_text=full_text)

    boundaries = {0, len(full_text)}
    for span in spans:
        boundaries.add(span.start)
        boundaries.add(span.end)
    for finding in merged:
        boundaries.add(finding.start)
        boundaries.add(finding.end)

    cuts = sorted(boundaries)
    fragments: list[tuple[Any, str]] = []

    for a, b in zip(cuts, cuts[1:]):
        if a == b:
            continue

        owner = _run_span_for_offset(spans, a)
        covering = next((f for f in merged if f.start <= a and b <= f.end), None)

        if covering is None:
            text = full_text[a:b]
        else:
            masked_full = _same_length_mask_value(covering.label, covering.quote)
            rel_start = a - covering.start
            rel_end = b - covering.start
            text = masked_full[rel_start:rel_end]

        fragments.append((owner.rpr_xml, text))

    _rewrite_paragraph_from_fragments(paragraph, fragments)


def _page_words_with_offsets(page: fitz.Page) -> tuple[str, list[PDFWord]]:
    raw = page.get_text("words", sort=True)
    words: list[PDFWord] = []
    parts: list[str] = []
    cursor = 0
    prev_line: tuple[int, int] | None = None

    for x0, y0, x1, y1, token, block_no, line_no, word_no in raw:
        token = (token or "").strip()
        if not token:
            continue

        current_line = (block_no, line_no)
        if parts:
            separator = "\n" if prev_line != current_line else " "
            parts.append(separator)
            cursor += len(separator)

        start = cursor
        parts.append(token)
        cursor += len(token)
        end = cursor

        words.append(
            PDFWord(
                start=start,
                end=end,
                text=token,
                rect=fitz.Rect(x0, y0, x1, y1),
                block_no=block_no,
                line_no=line_no,
                word_no=word_no,
            )
        )
        prev_line = current_line

    return "".join(parts), words


def _pdf_rects_for_findings(
    findings: Sequence[TextFinding],
    words: Sequence[PDFWord],
    *,
    original_text: str,
) -> list[fitz.Rect]:
    rects: list[fitz.Rect] = []
    merged = merge_overlapping_findings(findings, original_text=original_text)

    for finding in merged:
        matched = [w for w in words if not (w.end <= finding.start or w.start >= finding.end)]
        if not matched:
            continue

        by_line: dict[tuple[int, int], list[PDFWord]] = {}
        for word in matched:
            by_line.setdefault((word.block_no, word.line_no), []).append(word)

        for line_words in by_line.values():
            rect = line_words[0].rect
            for word in line_words[1:]:
                rect = rect | word.rect
            rects.append(rect)

    unique: list[fitz.Rect] = []
    seen: set[tuple[float, float, float, float]] = set()
    for rect in rects:
        key = (round(rect.x0, 3), round(rect.y0, 3), round(rect.x1, 3), round(rect.y1, 3))
        if key not in seen:
            seen.add(key)
            unique.append(rect)
    return unique


def _masked_rect_specs(
    findings: Sequence[TextFinding],
    words: Sequence[PDFWord],
    *,
    original_text: str,
) -> list[tuple[fitz.Rect, str]]:
    specs: list[tuple[fitz.Rect, str]] = []
    merged = merge_overlapping_findings(findings, original_text=original_text)

    for finding in merged:
        matched = [w for w in words if not (w.end <= finding.start or w.start >= finding.end)]
        if not matched:
            continue

        lines = {(w.block_no, w.line_no) for w in matched}
        if len(lines) != 1:
            continue

        rect = matched[0].rect
        for word in matched[1:]:
            rect = rect | word.rect

        specs.append((rect, mask_value_by_target(finding.label, finding.quote)))

    return specs


def _apply_page_redactions(page: fitz.Page, *, image_mode: str) -> None:
    kwargs: dict[str, Any] = {}

    if image_mode == "none" and hasattr(fitz, "PDF_REDACT_IMAGE_NONE"):
        kwargs["images"] = fitz.PDF_REDACT_IMAGE_NONE
    elif image_mode == "pixels" and hasattr(fitz, "PDF_REDACT_IMAGE_PIXELS"):
        kwargs["images"] = fitz.PDF_REDACT_IMAGE_PIXELS

    if hasattr(fitz, "PDF_REDACT_LINE_ART_NONE"):
        kwargs["graphics"] = fitz.PDF_REDACT_LINE_ART_NONE
    if hasattr(fitz, "PDF_REDACT_TEXT_REMOVE"):
        kwargs["text"] = fitz.PDF_REDACT_TEXT_REMOVE

    page.apply_redactions(**kwargs)


def _mask_docx(
    *,
    source_path: Path,
    output_path: Path,
    sdp: GoogleSDPClient,
    payload: DataMaskingRequest,
    custom_redactions: Optional[Sequence[str]] = None,
) -> None:
    document = docx.Document(source_path)
    for paragraph in _iter_document_paragraphs(document):
        full_text, _ = _paragraph_run_spans(paragraph)
        if not full_text.strip():
            continue

        findings = _mask_findings_for_text(
            sdp=sdp,
            text=full_text,
            payload=payload,
            custom_redactions=custom_redactions,
        )

        if findings:
            mask_paragraph_runs(paragraph, findings)

    document.save(output_path)


def _mask_image_file(
    *,
    source_path: Path,
    output_path: Path,
    sdp: GoogleSDPClient,
    payload: DataMaskingRequest,
    ocr_languages: Optional[Sequence[str]] = None,
    custom_redactions: Optional[Sequence[str]] = None,
) -> None:
    image = Image.open(source_path).convert("RGB")
    page_text, words = _ocr_words_from_image(image, ocr_lang=resolve_ocr_lang(ocr_languages))
    findings = _mask_findings_for_text(
        sdp=sdp,
        text=page_text,
        payload=payload,
        custom_redactions=custom_redactions,
    )
    boxes = _boxes_for_text_spans(findings, words)
    result = pixelate_boxes(image, boxes)
    fmt = DocumentInputFormat(source_path.suffix.lower().lstrip("."))
    result.save(output_path, format=_IMAGE_OUTPUT_MAP[fmt])


def _mask_pdf(
    *,
    source_path: Path,
    output_path: Path,
    sdp: GoogleSDPClient,
    payload: DataMaskingRequest,
    ocr_languages: Optional[Sequence[str]] = None,
    render_scale: float = DEFAULT_PDF_RENDER_SCALE,
    custom_redactions: Optional[Sequence[str]] = None,
) -> None:
    doc = fitz.open(source_path)
    try:
        for page in doc:
            page_text, words = _page_words_with_offsets(page)

            if words:
                findings = _mask_findings_for_text(
                    sdp=sdp,
                    text=page_text,
                    payload=payload,
                    custom_redactions=custom_redactions,
                )

                specs = _masked_rect_specs(findings, words, original_text=page_text)
                used_rects: set[tuple[float, float, float, float]] = set()

                for rect, masked in specs:
                    key = (round(rect.x0, 3), round(rect.y0, 3), round(rect.x1, 3), round(rect.y1, 3))
                    used_rects.add(key)
                    page.add_redact_annot(
                        rect,
                        text=masked,
                        fontname="helv",
                        fontsize=8,
                        fill=(1, 1, 1),
                        text_color=(0, 0, 0),
                        cross_out=False,
                    )

                for rect in _pdf_rects_for_findings(findings, words, original_text=page_text):
                    key = (round(rect.x0, 3), round(rect.y0, 3), round(rect.x1, 3), round(rect.y1, 3))
                    if key not in used_rects:
                        page.add_redact_annot(rect, fill=(1, 1, 1), cross_out=False)

                _apply_page_redactions(page, image_mode="none")
            else:
                pix = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale))
                image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                page_text, ocr_words = _ocr_words_from_image(image, ocr_lang=resolve_ocr_lang(ocr_languages))
                findings = _mask_findings_for_text(
                    sdp=sdp,
                    text=page_text,
                    payload=payload,
                    custom_redactions=custom_redactions,
                )
                boxes = _boxes_for_text_spans(findings, ocr_words)

                for x0, y0, x1, y1 in boxes:
                    scaled = fitz.Rect(
                        x0 / render_scale,
                        y0 / render_scale,
                        x1 / render_scale,
                        y1 / render_scale,
                    )
                    page.add_redact_annot(scaled, fill=(1, 1, 1), cross_out=False)

                _apply_page_redactions(page, image_mode="pixels")

        doc.save(output_path, garbage=4, deflate=True, clean=True)
    finally:
        doc.close()


def apply_data_mask(
    request: AnalyzerRequest | Mapping[str, Any],
    *,
    source_path: str | Path,
    output_dir: str | Path,
    sdp: GoogleSDPClient | None = None,
    project_id: str | None = None,
    location: str = DEFAULT_DLP_LOCATION,
    min_likelihood: str = DEFAULT_MIN_LIKELIHOOD,
    client: Any | None = None,
    ocr_languages: Optional[Sequence[str]] = None,
    custom_redactions: Optional[Sequence[str]] = None,
) -> AnalyzerResponse:
    req = validate_analyzer_request(request)
    input_payload, payload = _ensure_mask_request(req)
    source = Path(source_path)
    if not source.exists():
        raise FileNotFoundError(f"Source file not found: {source}")
    output_path = _output_path(source, Path(output_dir), suffix="_masked")
    resolved = _resolve_sdp(
        sdp=sdp,
        project_id=project_id,
        location=location,
        min_likelihood=min_likelihood,
        client=client,
    )

    fmt = input_payload.metadata.input_format
    if fmt == DocumentInputFormat.docx:
        _mask_docx(
            source_path=source,
            output_path=output_path,
            sdp=resolved,
            payload=payload,
            custom_redactions=custom_redactions,
        )
    elif fmt == DocumentInputFormat.pdf:
        _mask_pdf(
            source_path=source,
            output_path=output_path,
            sdp=resolved,
            payload=payload,
            ocr_languages=ocr_languages,
            custom_redactions=custom_redactions,
        )
    elif fmt in {DocumentInputFormat.jpg, DocumentInputFormat.jpeg, DocumentInputFormat.png}:
        _mask_image_file(
            source_path=source,
            output_path=output_path,
            sdp=resolved,
            payload=payload,
            ocr_languages=ocr_languages,
            custom_redactions=custom_redactions,
        )
    else:
        raise ValueError("Unsupported data_mask input format.")
    return _build_masking_response(request=req, output_path=output_path)


__all__ = [
    "preview_data_mask_candidates",
    "mask_value_by_target",
    "mask_text_with_findings",
    "pixelate_boxes",
    "mask_paragraph_runs",
    "apply_data_mask",
]
