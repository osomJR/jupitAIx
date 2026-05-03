
from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Mapping, Optional, Sequence

import docx
import fitz  # PyMuPDF
import pytesseract
from PIL import Image, ImageDraw

from src.extraction import OCR_CONFIG, preprocess_for_ocr, resolve_ocr_lang
from src.processing.data_protection.client import (
    DEFAULT_DLP_LOCATION,
    DEFAULT_MIN_LIKELIHOOD,
    DetectionCandidate,
    GoogleSDPClient,
    TextFinding,
    build_google_sdp_client,
    inspect_sensitive_text,
    merge_overlapping_findings,
    preview_candidates_from_text,
)
from src.schema import (
    AnalyzerRequest,
    AnalyzerResponse,
    DocumentInputFormat,
    DocumentPayload,
    FeatureType,
    RedactionRequest,
)
from src.validation import build_document_file_result, validate_analyzer_request, validate_analyzer_response

DEFAULT_PDF_RENDER_SCALE = 2.0
DEFAULT_BLACK = (0, 0, 0)

_IMAGE_OUTPUT_MAP: dict[DocumentInputFormat, str] = {
    DocumentInputFormat.jpg: "JPEG",
    DocumentInputFormat.jpeg: "JPEG",
    DocumentInputFormat.png: "PNG",
}

_FILE_OUTPUT_MAP = {
    DocumentInputFormat.pdf: "pdf",
    DocumentInputFormat.docx: "docx",
    DocumentInputFormat.jpg: "jpg",
    DocumentInputFormat.jpeg: "jpeg",
    DocumentInputFormat.png: "png",
}

_SUPPORTED_INPUTS = {
    DocumentInputFormat.pdf,
    DocumentInputFormat.docx,
    DocumentInputFormat.jpg,
    DocumentInputFormat.jpeg,
    DocumentInputFormat.png,
}


@dataclass(frozen=True)
class OCRWord:
    text: str
    bbox: tuple[int, int, int, int]
    start: int
    end: int


@dataclass(frozen=True)
class RunSpan:
    start: int
    end: int
    text: str
    rpr_xml: Any


@dataclass(frozen=True)
class PDFWord:
    start: int
    end: int
    text: str
    rect: fitz.Rect
    block_no: int
    line_no: int
    word_no: int


def redact_text_with_findings(text: str, findings: Sequence[TextFinding]) -> str:
    result = text
    for finding in sorted(findings, key=lambda item: item.start, reverse=True):
        replacement = "█" * max(1, finding.end - finding.start)
        result = result[:finding.start] + replacement + result[finding.end:]
    return result


def _ocr_words_from_image(image: Image.Image, *, ocr_lang: Optional[str] = None) -> tuple[str, list[OCRWord]]:
    processed = preprocess_for_ocr(image.convert("RGB"))
    data = pytesseract.image_to_data(
        processed,
        lang=ocr_lang or resolve_ocr_lang(),
        config=OCR_CONFIG,
        output_type=pytesseract.Output.DICT,
    )

    words: list[OCRWord] = []
    text_parts: list[str] = []
    cursor = 0
    texts = data.get("text", [])
    lefts = data.get("left", [])
    tops = data.get("top", [])
    widths = data.get("width", [])
    heights = data.get("height", [])

    for idx, raw in enumerate(texts):
        token = (raw or "").strip()
        if not token:
            continue
        if text_parts:
            text_parts.append(" ")
            cursor += 1
        start = cursor
        text_parts.append(token)
        cursor += len(token)
        end = cursor
        bbox = (
            int(lefts[idx]),
            int(tops[idx]),
            int(lefts[idx]) + int(widths[idx]),
            int(tops[idx]) + int(heights[idx]),
        )
        words.append(OCRWord(text=token, bbox=bbox, start=start, end=end))
    return "".join(text_parts), words


def _normalize_text_for_compare(value: str) -> str:
    import re

    return re.sub(r"\s+", " ", value.strip()).casefold()


def _boxes_for_text_spans(findings: Sequence[TextFinding], words: Sequence[OCRWord]) -> list[tuple[int, int, int, int]]:
    boxes: list[tuple[int, int, int, int]] = []
    for finding in findings:
        matched = [w for w in words if not (w.end <= finding.start or w.start >= finding.end)]
        if not matched:
            needle = _normalize_text_for_compare(finding.quote)
            for i in range(len(words)):
                buf = []
                for j in range(i, min(len(words), i + 12)):
                    buf.append(words[j])
                    candidate = _normalize_text_for_compare(" ".join(item.text for item in buf))
                    if candidate == needle:
                        matched = buf
                        break
                if matched:
                    break
        if not matched:
            continue
        x0 = min(w.bbox[0] for w in matched)
        y0 = min(w.bbox[1] for w in matched)
        x1 = max(w.bbox[2] for w in matched)
        y1 = max(w.bbox[3] for w in matched)
        boxes.append((x0, y0, x1, y1))

    unique: list[tuple[int, int, int, int]] = []
    seen: set[tuple[int, int, int, int]] = set()
    for box in boxes:
        if box not in seen:
            unique.append(box)
            seen.add(box)
    return unique


def draw_redaction_boxes(image: Image.Image, boxes: Sequence[tuple[int, int, int, int]]) -> Image.Image:
    out = image.convert("RGB").copy()
    draw = ImageDraw.Draw(out)
    for box in boxes:
        draw.rectangle(box, fill=DEFAULT_BLACK)
    return out


def _iter_table_paragraphs(table: Any):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_document_paragraphs(document: docx.Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _ensure_redaction_request(request: AnalyzerRequest) -> tuple[DocumentPayload, RedactionRequest]:
    if request.action != FeatureType.redact:
        raise ValueError("redact.py only handles action='redact'.")
    if not isinstance(request.input, DocumentPayload):
        raise ValueError("redact requires DocumentPayload input.")
    if not isinstance(request.payload, RedactionRequest):
        raise ValueError("redact requires RedactionRequest payload.")
    if request.input.metadata.input_format not in _SUPPORTED_INPUTS:
        raise ValueError("redact only supports pdf, docx, jpg, jpeg, png.")
    if request.policy.structure_preservation is not True:
        raise ValueError("structure_preservation must be True for redact.")
    return request.input, request.payload


def _output_path(source_path: Path, output_dir: Path, suffix: str = "_redacted") -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{source_path.stem}{suffix}{source_path.suffix.lower()}"


def _build_redaction_response(*, request: AnalyzerRequest, output_path: Path) -> AnalyzerResponse:
    input_payload, _ = _ensure_redaction_request(request)
    result = build_document_file_result(
        filename=output_path.name,
        output_format=_FILE_OUTPUT_MAP[input_payload.metadata.input_format],
        file_size_mb=round(output_path.stat().st_size / (1024 * 1024), 4),
        algorithm_version="google-sdp-redact-v1",
    )
    response = AnalyzerResponse(
        action=FeatureType.redact,
        input_format=input_payload.metadata.input_format,
        policy=request.policy,
        system_language=request.system_language,
        result=result,
    )
    return validate_analyzer_response(response, request=request)


def _resolve_sdp(
    *,
    sdp: GoogleSDPClient | None = None,
    project_id: str | None = None,
    location: str = DEFAULT_DLP_LOCATION,
    min_likelihood: str = DEFAULT_MIN_LIKELIHOOD,
    client: Any | None = None,
) -> GoogleSDPClient:
    if sdp is not None:
        return sdp
    if not project_id:
        raise ValueError("Provide either sdp or project_id.")
    return build_google_sdp_client(
        project_id=project_id,
        location=location,
        min_likelihood=min_likelihood,
        client=client,
    )


def preview_redaction_candidates(
    request: AnalyzerRequest | Mapping[str, Any],
    *,
    sdp: GoogleSDPClient | None = None,
    project_id: str | None = None,
    location: str = DEFAULT_DLP_LOCATION,
    min_likelihood: str = DEFAULT_MIN_LIKELIHOOD,
    client: Any | None = None,
) -> list[DetectionCandidate]:
    req = validate_analyzer_request(request)
    input_payload, payload = _ensure_redaction_request(req)
    if not input_payload.text:
        raise ValueError(
            "preview_redaction_candidates requires extracted document text. "
            "Build the request with extraction.py before previewing."
        )
    resolved = _resolve_sdp(
        sdp=sdp,
        project_id=project_id,
        location=location,
        min_likelihood=min_likelihood,
        client=client,
    )
    return preview_candidates_from_text(
        sdp=resolved,
        text=input_payload.text,
        targets=payload.target_data,
        review_exclusions=payload.review_exclusions,
        min_likelihood=min_likelihood,
    )


def _paragraph_run_spans(paragraph: Any) -> tuple[str, list[RunSpan]]:
    spans: list[RunSpan] = []
    parts: list[str] = []
    cursor = 0

    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        start = cursor
        end = start + len(text)
        parts.append(text)
        spans.append(
            RunSpan(
                start=start,
                end=end,
                text=text,
                rpr_xml=deepcopy(run._r.rPr) if run._r.rPr is not None else None,
            )
        )
        cursor = end

    return "".join(parts), spans


def _copy_rpr(dst_run: Any, rpr_xml: Any) -> None:
    if dst_run._r.rPr is not None:
        dst_run._r.remove(dst_run._r.rPr)
    if rpr_xml is not None:
        dst_run._r.insert(0, deepcopy(rpr_xml))


def _rewrite_paragraph_from_fragments(paragraph: Any, fragments: list[tuple[Any, str]]) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    for rpr_xml, text in fragments:
        if not text:
            continue
        new_run = paragraph.add_run()
        _copy_rpr(new_run, rpr_xml)
        new_run.text = text


def _run_span_for_offset(spans: Sequence[RunSpan], offset: int) -> RunSpan:
    for span in spans:
        if span.start <= offset < span.end:
            return span
    return spans[-1]


def redact_paragraph_runs(paragraph: Any, findings: Sequence[TextFinding]) -> None:
    full_text, spans = _paragraph_run_spans(paragraph)
    if not full_text or not spans or not findings:
        return

    merged = merge_overlapping_findings(findings, original_text=full_text)

    boundaries = {0, len(full_text)}
    for span in spans:
        boundaries.add(span.start)
        boundaries.add(span.end)
    for finding in merged:
        boundaries.add(finding.start)
        boundaries.add(finding.end)

    cuts = sorted(boundaries)
    fragments: list[tuple[Any, str]] = []

    for a, b in zip(cuts, cuts[1:]):
        if a == b:
            continue

        owner = _run_span_for_offset(spans, a)
        covered = any(f.start <= a and b <= f.end for f in merged)

        if covered:
            text = "█" * (b - a)
        else:
            text = full_text[a:b]

        fragments.append((owner.rpr_xml, text))

    _rewrite_paragraph_from_fragments(paragraph, fragments)


def _page_words_with_offsets(page: fitz.Page) -> tuple[str, list[PDFWord]]:
    raw = page.get_text("words", sort=True)
    words: list[PDFWord] = []
    parts: list[str] = []
    cursor = 0
    prev_line: tuple[int, int] | None = None

    for x0, y0, x1, y1, token, block_no, line_no, word_no in raw:
        token = (token or "").strip()
        if not token:
            continue

        current_line = (block_no, line_no)
        if parts:
            separator = "\n" if prev_line != current_line else " "
            parts.append(separator)
            cursor += len(separator)

        start = cursor
        parts.append(token)
        cursor += len(token)
        end = cursor

        words.append(
            PDFWord(
                start=start,
                end=end,
                text=token,
                rect=fitz.Rect(x0, y0, x1, y1),
                block_no=block_no,
                line_no=line_no,
                word_no=word_no,
            )
        )
        prev_line = current_line

    return "".join(parts), words


def _pdf_rects_for_findings(
    findings: Sequence[TextFinding],
    words: Sequence[PDFWord],
    *,
    original_text: str,
) -> list[fitz.Rect]:
    rects: list[fitz.Rect] = []
    merged = merge_overlapping_findings(findings, original_text=original_text)

    for finding in merged:
        matched = [w for w in words if not (w.end <= finding.start or w.start >= finding.end)]
        if not matched:
            continue

        by_line: dict[tuple[int, int], list[PDFWord]] = {}
        for word in matched:
            by_line.setdefault((word.block_no, word.line_no), []).append(word)

        for line_words in by_line.values():
            rect = line_words[0].rect
            for word in line_words[1:]:
                rect = rect | word.rect
            rects.append(rect)

    unique: list[fitz.Rect] = []
    seen: set[tuple[float, float, float, float]] = set()
    for rect in rects:
        key = (round(rect.x0, 3), round(rect.y0, 3), round(rect.x1, 3), round(rect.y1, 3))
        if key not in seen:
            seen.add(key)
            unique.append(rect)
    return unique


def _apply_page_redactions(page: fitz.Page, *, image_mode: str) -> None:
    kwargs: dict[str, Any] = {}

    if image_mode == "none" and hasattr(fitz, "PDF_REDACT_IMAGE_NONE"):
        kwargs["images"] = fitz.PDF_REDACT_IMAGE_NONE
    elif image_mode == "pixels" and hasattr(fitz, "PDF_REDACT_IMAGE_PIXELS"):
        kwargs["images"] = fitz.PDF_REDACT_IMAGE_PIXELS

    if hasattr(fitz, "PDF_REDACT_LINE_ART_NONE"):
        kwargs["graphics"] = fitz.PDF_REDACT_LINE_ART_NONE
    if hasattr(fitz, "PDF_REDACT_TEXT_REMOVE"):
        kwargs["text"] = fitz.PDF_REDACT_TEXT_REMOVE

    page.apply_redactions(**kwargs)


def _redact_docx(
    *,
    source_path: Path,
    output_path: Path,
    sdp: GoogleSDPClient,
    payload: RedactionRequest,
) -> None:
    document = docx.Document(source_path)
    for paragraph in _iter_document_paragraphs(document):
        full_text, _ = _paragraph_run_spans(paragraph)
        if not full_text.strip():
            continue

        findings = inspect_sensitive_text(
            sdp=sdp,
            text=full_text,
            targets=payload.target_data,
            review_exclusions=payload.review_exclusions,
        )

        if findings:
            redact_paragraph_runs(paragraph, findings)

    document.save(output_path)


def _redact_image_file(
    *,
    source_path: Path,
    output_path: Path,
    sdp: GoogleSDPClient,
    payload: RedactionRequest,
    ocr_languages: Optional[Sequence[str]] = None,
) -> None:
    image = Image.open(source_path).convert("RGB")
    page_text, words = _ocr_words_from_image(image, ocr_lang=resolve_ocr_lang(ocr_languages))
    findings = inspect_sensitive_text(
        sdp=sdp,
        text=page_text,
        targets=payload.target_data,
        review_exclusions=payload.review_exclusions,
    )
    boxes = _boxes_for_text_spans(findings, words)
    result = draw_redaction_boxes(image, boxes)
    fmt = DocumentInputFormat(source_path.suffix.lower().lstrip("."))
    result.save(output_path, format=_IMAGE_OUTPUT_MAP[fmt])


def _redact_pdf(
    *,
    source_path: Path,
    output_path: Path,
    sdp: GoogleSDPClient,
    payload: RedactionRequest,
    ocr_languages: Optional[Sequence[str]] = None,
    render_scale: float = DEFAULT_PDF_RENDER_SCALE,
) -> None:
    doc = fitz.open(source_path)
    try:
        for page in doc:
            page_text, words = _page_words_with_offsets(page)

            if words:
                findings = inspect_sensitive_text(
                    sdp=sdp,
                    text=page_text,
                    targets=payload.target_data,
                    review_exclusions=payload.review_exclusions,
                )

                for rect in _pdf_rects_for_findings(findings, words, original_text=page_text):
                    page.add_redact_annot(rect, fill=DEFAULT_BLACK, cross_out=False)

                _apply_page_redactions(page, image_mode="none")
            else:
                pix = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale))
                image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                page_text, ocr_words = _ocr_words_from_image(image, ocr_lang=resolve_ocr_lang(ocr_languages))
                findings = inspect_sensitive_text(
                    sdp=sdp,
                    text=page_text,
                    targets=payload.target_data,
                    review_exclusions=payload.review_exclusions,
                )
                boxes = _boxes_for_text_spans(findings, ocr_words)

                for x0, y0, x1, y1 in boxes:
                    scaled = fitz.Rect(
                        x0 / render_scale,
                        y0 / render_scale,
                        x1 / render_scale,
                        y1 / render_scale,
                    )
                    page.add_redact_annot(scaled, fill=DEFAULT_BLACK, cross_out=False)

                _apply_page_redactions(page, image_mode="pixels")

        doc.save(output_path, garbage=4, deflate=True, clean=True)
    finally:
        doc.close()


def apply_redaction(
    request: AnalyzerRequest | Mapping[str, Any],
    *,
    source_path: str | Path,
    output_dir: str | Path,
    sdp: GoogleSDPClient | None = None,
    project_id: str | None = None,
    location: str = DEFAULT_DLP_LOCATION,
    min_likelihood: str = DEFAULT_MIN_LIKELIHOOD,
    client: Any | None = None,
    ocr_languages: Optional[Sequence[str]] = None,
) -> AnalyzerResponse:
    req = validate_analyzer_request(request)
    input_payload, payload = _ensure_redaction_request(req)
    source = Path(source_path)
    if not source.exists():
        raise FileNotFoundError(f"Source file not found: {source}")
    output_path = _output_path(source, Path(output_dir), suffix="_redacted")
    resolved = _resolve_sdp(
        sdp=sdp,
        project_id=project_id,
        location=location,
        min_likelihood=min_likelihood,
        client=client,
    )

    fmt = input_payload.metadata.input_format
    if fmt == DocumentInputFormat.docx:
        _redact_docx(source_path=source, output_path=output_path, sdp=resolved, payload=payload)
    elif fmt == DocumentInputFormat.pdf:
        _redact_pdf(
            source_path=source,
            output_path=output_path,
            sdp=resolved,
            payload=payload,
            ocr_languages=ocr_languages,
        )
    elif fmt in {DocumentInputFormat.jpg, DocumentInputFormat.jpeg, DocumentInputFormat.png}:
        _redact_image_file(
            source_path=source,
            output_path=output_path,
            sdp=resolved,
            payload=payload,
            ocr_languages=ocr_languages,
        )
    else:
        raise ValueError("Unsupported redact input format.")
    return _build_redaction_response(request=req, output_path=output_path)


__all__ = [
    "preview_redaction_candidates",
    "redact_text_with_findings",
    "draw_redaction_boxes",
    "redact_paragraph_runs",
    "apply_redaction",
]
